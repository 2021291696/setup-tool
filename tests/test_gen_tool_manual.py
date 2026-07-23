"""Regression tests for concise, evidence-backed manuals."""
import json
import sys
from zipfile import ZipFile
from pathlib import Path

import pytest

sys.path.insert(0, str(Path(__file__).resolve().parent.parent / "scripts"))
import gen_tool_manual as gtm  # noqa: E402


SAMPLE = {
    "name": "test-tool",
    "function_desc": "Automates a focused, repetitive task.",
    "usage": ["Run the documented command.", "Check the generated result."],
    "special_notes": ["Requires noncommercial use."],
    "sources": ["README: https://example.com/test-tool#usage"],
}


def test_manual_contains_only_reader_facing_sections(tmp_path):
    content = gtm.generate_markdown([SAMPLE], tmp_path).read_text(encoding="utf-8")
    assert content.startswith("# test-tool")
    for section in ("## 作用", "## 使用方法", "## 特殊注意事项"):
        assert section in content
    for forbidden in ("安装方案", "最终选择", "验证结果", "来源与核验", "适配性评估", "配置"):
        assert forbidden not in content


def test_special_notes_section_is_omitted_when_empty(tmp_path):
    tool = {**SAMPLE, "special_notes": []}
    content = gtm.generate_markdown([tool], tmp_path).read_text(encoding="utf-8")
    assert "特殊注意事项" not in content


def test_usage_is_never_derived_from_name(tmp_path):
    tool = {**SAMPLE, "name": "no-trigger-tool", "usage": ["Open it from the documented host menu."]}
    content = gtm.generate_markdown([tool], tmp_path).read_text(encoding="utf-8")
    assert "/no-trigger-tool" not in content
    assert "Open it from the documented host menu." in content


@pytest.mark.parametrize("field", ["function_desc", "usage", "sources"])
def test_missing_required_metadata_fails_fast(tmp_path, field):
    tool = dict(SAMPLE)
    tool.pop(field)
    with pytest.raises(ValueError, match=field):
        gtm.generate_markdown([tool], tmp_path)


def test_word_has_only_allowed_sections(tmp_path):
    output = gtm.generate_docx([SAMPLE], tmp_path)
    from docx import Document

    manual = Document(output)
    paragraphs = manual.paragraphs
    text = "\n".join(paragraph.text for paragraph in paragraphs)
    headings = [paragraph.text for paragraph in paragraphs if paragraph.style.name.startswith("Heading")]
    assert headings == ["test-tool", "作用", "使用方法", "特殊注意事项"]
    assert [paragraph.style.name for paragraph in paragraphs if paragraph.text in SAMPLE["usage"]] == [
        "List Number",
        "List Number",
    ]
    assert [paragraph.style.name for paragraph in paragraphs if paragraph.text in SAMPLE["special_notes"]] == [
        "List Bullet"
    ]
    assert all(section in text for section in ("作用", "使用方法", "特殊注意事项"))
    for forbidden in ("安装方案", "最终选择", "验证结果", "来源与核验", "适配性评估"):
        assert forbidden not in text
    assert SAMPLE["sources"][0] not in text
    with ZipFile(output) as archive:
        assert archive.testzip() is None
        assert b'w:percent="100"' in archive.read("word/settings.xml")
    assert round(manual.sections[0].page_width.mm) == 210
    assert round(manual.sections[0].page_height.mm) == 297


@pytest.mark.parametrize("special_notes", [[], None])
def test_word_omits_special_notes_section_when_empty(tmp_path, special_notes):
    tool = dict(SAMPLE)
    if special_notes is None:
        tool.pop("special_notes")
    else:
        tool["special_notes"] = special_notes
    output = gtm.generate_docx([tool], tmp_path)
    from docx import Document

    text = "\n".join(paragraph.text for paragraph in Document(output).paragraphs)
    assert "特殊注意事项" not in text


def test_cli_defaults_to_word_only(tmp_path, monkeypatch):
    metadata = tmp_path / "metadata.json"
    metadata.write_text(json.dumps(SAMPLE, ensure_ascii=False), encoding="utf-8")
    monkeypatch.setattr(
        sys,
        "argv",
        ["gen_tool_manual.py", "--file", str(metadata), "--out-dir", str(tmp_path)],
    )

    assert gtm.main() == 0

    assert (tmp_path / "test-tool.docx").is_file()
    assert not (tmp_path / "test-tool.md").exists()


@pytest.mark.parametrize(
    ("manual_format", "expected_suffixes"),
    [("markdown", {".md"}), ("both", {".md", ".docx"})],
)
def test_cli_keeps_explicit_format_options(tmp_path, monkeypatch, manual_format, expected_suffixes):
    metadata = tmp_path / "metadata.json"
    metadata.write_text(json.dumps(SAMPLE, ensure_ascii=False), encoding="utf-8")
    monkeypatch.setattr(
        sys,
        "argv",
        [
            "gen_tool_manual.py",
            "--file",
            str(metadata),
            "--out-dir",
            str(tmp_path),
            "--format",
            manual_format,
        ],
    )

    assert gtm.main() == 0
    actual_suffixes = {
        path.suffix for path in tmp_path.glob("test-tool.*") if path.suffix in {".md", ".docx"}
    }
    assert actual_suffixes == expected_suffixes


@pytest.mark.parametrize("name", ["CON", "nul.txt", "bad.", "bad\x01name"])
def test_windows_unsafe_filenames_are_rejected(tmp_path, name):
    with pytest.raises(ValueError, match="safe filename"):
        gtm.generate_markdown([{**SAMPLE, "name": name}], tmp_path)


def test_utf8_output_configuration_is_safe():
    gtm.configure_utf8_output()
