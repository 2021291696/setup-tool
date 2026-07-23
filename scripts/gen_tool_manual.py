"""Generate concise, evidence-backed tool manuals.

The reader-facing manual contains only purpose, usage, and optional special
notes. Sources remain required metadata for correctness but are not rendered.
Word is the default output; Markdown is optional.
"""
import argparse
import json
import sys
from pathlib import Path
from zipfile import BadZipFile, ZipFile

OUT_DIR = Path("工具说明书")
DEFAULT_FORMAT = "docx"
WINDOWS_RESERVED_NAMES = {
    "CON",
    "PRN",
    "AUX",
    "NUL",
    *(f"COM{index}" for index in range(1, 10)),
    *(f"LPT{index}" for index in range(1, 10)),
}


def configure_utf8_output() -> None:
    """Keep paths readable when Windows captures stdout through a pipe."""
    for stream in (sys.stdout, sys.stderr):
        if hasattr(stream, "reconfigure"):
            stream.reconfigure(encoding="utf-8")


def _require_text(data: dict, key: str) -> str:
    value = data.get(key)
    if not isinstance(value, str) or not value.strip():
        raise ValueError(f"{key} must be a non-empty string")
    return value.strip()


def _string_list(value: object, field: str, *, required: bool = False) -> list[str]:
    if value is None:
        value = []
    if not isinstance(value, list) or any(not isinstance(item, str) or not item.strip() for item in value):
        raise ValueError(f"{field} must be a list of non-empty strings")
    result = list(dict.fromkeys(item.strip() for item in value))
    if required and not result:
        raise ValueError(f"{field} must not be empty")
    return result


def validate_tool(tool: dict) -> dict:
    """Reject incomplete or unsourced usage before writing a manual."""
    if not isinstance(tool, dict):
        raise ValueError("each tool must be an object")
    name = _require_text(tool, "name")
    reserved_stem = name.split(".", 1)[0].upper()
    if (
        any(char in name for char in '\\\\/:*?\"<>|')
        or any(ord(char) < 32 for char in name)
        or name.endswith(".")
        or name in {".", ".."}
        or reserved_stem in WINDOWS_RESERVED_NAMES
    ):
        raise ValueError("name must be a safe filename")
    return {
        "name": name,
        "function_desc": _require_text(tool, "function_desc"),
        "usage": _string_list(tool.get("usage"), "usage", required=True),
        "special_notes": _string_list(tool.get("special_notes"), "special_notes"),
        "sources": _string_list(tool.get("sources"), "sources", required=True),
    }


def _bullet(lines: list[str], values: list[str]) -> None:
    lines.extend(f"{index}. {value}" for index, value in enumerate(values, 1))


def render_markdown(tool: dict, heading_level: int) -> str:
    """Render only the three allowed reader-facing sections."""
    section = "#" * heading_level
    subsection = "#" * (heading_level + 1)
    lines = [f"{section} {tool['name']}", "", f"{subsection} 作用", "", tool["function_desc"], "", f"{subsection} 使用方法", ""]
    _bullet(lines, tool["usage"])
    if tool["special_notes"]:
        lines.extend(["", f"{subsection} 特殊注意事项", ""])
        lines.extend(f"- {note}" for note in tool["special_notes"])
    return "\n".join(lines)


def build_filename(tools: list, suffix: str) -> str:
    return "-".join(tool["name"] for tool in tools) + suffix


def _validated(tools: list) -> list[dict]:
    if not isinstance(tools, list) or not tools:
        raise ValueError("tools must be a non-empty list")
    return [validate_tool(tool) for tool in tools]


def generate_markdown(tools: list, out_dir: Path = OUT_DIR) -> Path:
    tools = _validated(tools)
    out_dir.mkdir(parents=True, exist_ok=True)
    if len(tools) == 1:
        content = render_markdown(tools[0], 1) + "\n"
    else:
        title = " + ".join(tool["name"] for tool in tools)
        content = f"# {title} 使用说明\n\n" + "\n\n---\n\n".join(render_markdown(tool, 2) for tool in tools) + "\n"
    output = out_dir / build_filename(tools, ".md")
    output.write_text(content, encoding="utf-8")
    return output


def generate_docx(tools: list, out_dir: Path = OUT_DIR) -> Path:
    """Generate the default Word manual."""
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.shared import Mm, Pt, RGBColor
    except ImportError as exc:
        raise RuntimeError(
            "Word output requires python-docx; install it or explicitly choose --format markdown"
        ) from exc
    tools = _validated(tools)
    out_dir.mkdir(parents=True, exist_ok=True)
    document = Document()
    zoom = document.settings.element.find(qn("w:zoom"))
    if zoom is not None:
        zoom.set(qn("w:percent"), "100")
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(20)
    section.right_margin = Mm(22)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(22)
    for style_name, size, bold in (
        ("Normal", 11, False),
        ("Heading 1", 18, True),
        ("Heading 2", 14, True),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style._element.get_or_add_rPr().get_or_add_rFonts().set(
            qn("w:eastAsia"), "Microsoft YaHei"
        )
    for tool_index, tool in enumerate(tools):
        if tool_index:
            document.add_page_break()
        document.add_heading(tool["name"], level=1)
        document.add_heading("作用", level=2)
        document.add_paragraph(tool["function_desc"])
        document.add_heading("使用方法", level=2)
        for step in tool["usage"]:
            document.add_paragraph(step, style="List Number")
        if tool["special_notes"]:
            document.add_heading("特殊注意事项", level=2)
            for note in tool["special_notes"]:
                document.add_paragraph(note, style="List Bullet")
    output = out_dir / build_filename(tools, ".docx")
    document.save(output)
    _verify_docx(output, tools, Document)
    return output


def _verify_docx(output: Path, tools: list[dict], document_class: type) -> None:
    """Fail if the saved Word package cannot be read back as the intended manual."""
    if not output.is_file() or output.stat().st_size == 0:
        raise RuntimeError(f"Word output was not created correctly: {output}")
    try:
        with ZipFile(output) as archive:
            damaged_member = archive.testzip()
        if damaged_member is not None:
            raise RuntimeError(f"Word output contains a damaged member: {damaged_member}")
        paragraphs = document_class(output).paragraphs
    except (BadZipFile, OSError, ValueError) as exc:
        raise RuntimeError(f"Word output could not be read back: {output}") from exc

    expected_headings = []
    expected_body = []
    for tool in tools:
        expected_headings.extend((tool["name"], "作用", "使用方法"))
        expected_body.extend((tool["function_desc"], *tool["usage"]))
        if tool["special_notes"]:
            expected_headings.append("特殊注意事项")
            expected_body.extend(tool["special_notes"])

    actual_headings = [
        paragraph.text
        for paragraph in paragraphs
        if paragraph.style.style_id in {"Heading1", "Heading2"}
    ]
    if actual_headings != expected_headings:
        raise RuntimeError("Word output headings do not match the concise manual contract")
    paragraph_text = [paragraph.text for paragraph in paragraphs]
    if any(value not in paragraph_text for value in expected_body):
        raise RuntimeError("Word output is missing reader-facing content")
    rendered_text = "\n".join(paragraph_text)
    if any(source in rendered_text for tool in tools for source in tool["sources"]):
        raise RuntimeError("Word output leaked internal source metadata")


def build_parser() -> argparse.ArgumentParser:
    parser = argparse.ArgumentParser()
    parser.add_argument("--file", required=True, help="metadata JSON file")
    parser.add_argument("--out-dir", default=str(OUT_DIR))
    parser.add_argument("--format", choices=("docx", "markdown", "both"), default=DEFAULT_FORMAT)
    return parser


def main() -> int:
    configure_utf8_output()
    args = build_parser().parse_args()
    try:
        with open(args.file, encoding="utf-8") as source:
            tools = json.load(source)
        if isinstance(tools, dict):
            tools = [tools]
        outputs = []
        if args.format in ("markdown", "both"):
            outputs.append(generate_markdown(tools, Path(args.out_dir)))
        if args.format in ("docx", "both"):
            outputs.append(generate_docx(tools, Path(args.out_dir)))
    except (OSError, json.JSONDecodeError, ValueError, RuntimeError) as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    print("OK: wrote " + ", ".join(str(path) for path in outputs))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
